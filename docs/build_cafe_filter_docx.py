"""Generate a formatted Cafes filter .docx from the v2 spec.

Usage:
    pip install python-docx
    python docs/build_cafe_filter_docx.py

Output: docs/Cafe_Filter_Documentation.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


MONO = "Consolas"


def mono_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = MONO
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def kv_table(doc, rows, headers=("Setting", "Value")):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = headers
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
    doc.add_paragraph()
    return table


def grid_table(doc, header, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def main():
    out = Path(__file__).resolve().parent / "Cafe_Filter_Documentation.docx"
    doc = Document()

    title = doc.add_heading("Cafes POI Filter", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "dreamneighborhood.com - Overture Places filtering - "
        "Status: SHIPPABLE (v2, no excludes)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph(
        "Validated across 4 metros: Fort Pierce FL (22), Warren OH (35), "
        "San Jose CA (250+), Ely NV / Park City UT (5). "
        "Pure category gate, no name or category excludes. Gate on "
        "basic_category = coffee_shop, NOT cafe."
    )

    doc.add_heading("1. Production filter - copy/paste into the admin", level=1)

    kv_table(doc, [
        ("Saved filter", "Cafe"),
        ("Distance", "10.0 miles (1-3 urban, 10 rural)"),
        ("Max results", "250"),
        ("Min confidence", "(empty)"),
        ("Operating status", "open"),
        ("Deduplicate addresses", "ON"),
    ])

    doc.add_heading("Include", level=2)
    doc.add_paragraph("Basic category (include):")
    mono_block(doc, ["coffee_shop", "tea_room"])
    doc.add_paragraph("Category primary (include):")
    mono_block(doc, ["coffee_shop", "tea_room", "coffee_roastery"])
    doc.add_paragraph(
        "coffee_roastery is required: Headframe Coffee (Ely NV) has "
        "category_primary = coffee_roastery."
    )

    doc.add_heading("Exclude", level=2)
    doc.add_paragraph("Leave EVERY exclude box empty. See sections 3-4.")

    doc.add_heading("Query builder", level=2)
    mono_block(doc, ["basic_category_include or category_primary_include"])

    doc.add_heading("2. Why basic_category = cafe failed", level=1)
    grid_table(
        doc,
        ["Field", "Schema docs said", "Your data has"],
        [
            ["basic_category", "cafe", "coffee_shop"],
            ["taxonomy_hierarchy", "food_and_drink, cafe, coffee_shop",
             "food_and_drink, non_alcoholic_beverage_venue, coffee_shop"],
        ],
        widths=[1.5, 2.5, 3.5],
    )

    doc.add_heading("3. Why name excludes are banned (USA-wide)", level=1)
    doc.add_paragraph(
        "A name token safe in one town deletes real cafes elsewhere. There is no "
        "finite name list for 50 states."
    )
    grid_table(
        doc,
        ["Tempting token", "Meant to drop", "Also kills (real cafes)"],
        [
            ["*market*", "Streetside Market (FL)", "Corner Market Kitchen (SJ)"],
            ["Sunoco / Citgo / APlus", "gas-station coffee", "infinite, unwinnable list"],
            ["nutrition", "Powerful Nutrition club", "any 'Nutrition Cafe'"],
        ],
        widths=[1.8, 2.4, 3.0],
    )

    doc.add_heading("4. Why category excludes are ALSO out", level=1)
    doc.add_paragraph(
        "Junk alternates ride on legitimate coffee shops, so excluding them "
        "removes more signal than noise."
    )
    grid_table(
        doc,
        ["Alternate", "Junk removed", "Real cafes wrongly deleted"],
        [
            ["bubble_tea", "boba counters", "Nirvana Soul Coffee, Dr.ink, Cafe Boba"],
            ["smoothie_juice_bar", "Protein Harbor", "Voyager Craft Coffee HQ, Cowboy Coffee"],
            ["donuts / fast_food_restaurant", "-", "every Dunkin'"],
            ["restaurant", "-", "Starbucks, Peet's, most of the list"],
        ],
        widths=[1.9, 1.7, 3.6],
    )

    doc.add_heading("5. The one real knob: include tea_room or not", level=1)
    grid_table(
        doc,
        ["Option", "Includes tea_room?", "Effect"],
        [
            ["A - Coffee + Tea (default)", "Yes", "Broad; also pulls in boba counters"],
            ["B - Coffee-only", "No", "Drops boba wholesale; also drops real tea rooms"],
        ],
        widths=[2.3, 1.5, 3.4],
    )

    doc.add_heading("6. Residual noise - needs data enrichment", level=1)
    grid_table(
        doc,
        ["Problem", "Examples", "Needed flag"],
        [
            ["In-store counters", "WFM Coffee Bar, Nordstrom Ebar, Nespresso", "is_in_store_concession"],
            ["Gas-station coffee", "APlus at Sunoco", "is_fuel_station"],
            ["Drive-through-only", "7 Brew", "is_drive_through_only"],
            ["Nutrition clubs", "Powerful Nutrition", "is_nutrition_club"],
        ],
        widths=[1.8, 2.6, 2.8],
    )

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
