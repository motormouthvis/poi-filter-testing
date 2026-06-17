"""Generate a formatted Restaurants filter .docx from the v3 spec.

Usage:
    pip install python-docx
    python docs/build_restaurant_filter_docx.py

Output: docs/Restaurant_Filter_Documentation.docx
Upload to Google Drive, then Open with -> Google Docs to keep headings/tables.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


MONO = "Consolas"


def mono_block(doc, lines):
    """Render a list of strings as a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = MONO
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def kv_table(doc, rows, headers=("Setting", "Value")):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = headers
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
    doc.add_paragraph()
    return table


def grid_table(doc, header, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def main():
    out = Path(__file__).resolve().parent / "Restaurant_Filter_Documentation.docx"
    doc = Document()

    # ---- Title ----
    title = doc.add_heading("Restaurants POI Filter", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "dreamneighborhood.com - Overture Places filtering - Status: WORKING (v3)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph(
        "On the latest test (lat/lon 27.4612, -80.3035, 5 mi) this config returns "
        "104 results, down from 158 under v2. The worst noise is gone: the retail "
        "seafood market is dropped by the confidence gate, name-based market excludes "
        "catch grocery/fish/meat markets, and address dedup collapses true duplicates. "
        "Remaining issues are data-quality problems Overture cannot express in a boolean "
        "filter - see section 5 (Dev team: what to improve next)."
    )
    doc.add_paragraph(
        "Goal: surface real sit-down / casual restaurants for the mom-centric "
        "home-shopper widget, while excluding fast food, drinking-focused venues "
        "(bars, breweries), cafes, bakeries, food trucks, retail food markets, and "
        "services. This must be a generic USA solution - no hardcoded business names "
        "or locale-specific hacks."
    )
    doc.add_paragraph(
        "Match operators: token = contains, \"token\" = exact, token* = starts-with, "
        "*token = ends-with, *token* = contains-anywhere."
    )

    # ---- 1. Working filter ----
    doc.add_heading("1. The current working filter (v3) - copy/paste into the admin", level=1)

    doc.add_heading("Top controls", level=2)
    kv_table(doc, [
        ("Saved filter", "Restaurants"),
        ("Distance", "5.0 miles"),
        ("Max results", "250"),
        ("Min confidence", "0.7"),
        ("Operating status", "open"),
        ("Deduplicate addresses", "ON (keep highest-confidence record per address)"),
    ])

    doc.add_heading("Include", level=2)
    doc.add_paragraph("Basic category (include):")
    mono_block(doc, ["restaurant"])
    doc.add_paragraph(
        "Leave Category primary (include) and Category alternate (include) EMPTY. "
        "The basic_category = restaurant gate does all the inclusion; the old broad "
        "category includes were the original source of the non-restaurant leak."
    )

    doc.add_heading("Exclude", level=2)
    doc.add_paragraph("Business name primary (exclude):")
    mono_block(doc, [
        "taco truck", "food truck", "*market", "*market*", "supermarket",
        "grocery", "fish market", "meat market", "seafood market", "*tienda*",
    ])
    doc.add_paragraph("Category primary (exclude):")
    mono_block(doc, [
        "fast_food_restaurant", "food_truck", "street_vendor", "coffee_shop",
        "cafe", "bakery", "nightclub", "brewery", "distillery", "winery",
        "caterer", "food_court",
    ])
    doc.add_paragraph("Category alternate (exclude):")
    mono_block(doc, [
        "fast_food_restaurant", "food_truck", "coffee_shop", "cafe", "bakery",
        "nightclub", "brewery", "distillery", "winery", "caterer", "food_court",
    ])
    doc.add_paragraph(
        "Leave the remaining exclude boxes (name common, basic category, taxonomy "
        "primary/alternates) empty."
    )

    doc.add_heading("Query builder", level=2)
    mono_block(doc, [
        "basic_category_include and category_primary_exclude and "
        "category_alternate_exclude and name_primary_exclude",
    ])
    doc.add_paragraph(
        "Plain English: must be basic_category = restaurant, and not match any "
        "primary-category exclude, and not match any alternate-category exclude, and "
        "its primary name is not in the name-exclude list."
    )

    # ---- 2. v2 -> v3 ----
    doc.add_heading("2. What changed from v2 to v3 and why it helped", level=1)
    grid_table(
        doc,
        ["Lever", "v2", "v3", "Effect"],
        [
            ["Min confidence", "0.0", "0.7",
             "Drops the worst-tagged noise, incl. the Pelican Seafood retail market "
             "(conf 0.53) and ghost/duplicate fragments."],
            ["Deduplicate addresses", "OFF", "ON",
             "Collapses true duplicate rows (e.g. two Pot Belli Deli records, "
             "Goodfella's fragments at one address)."],
            ["Name excludes", "truck only", "+ market/grocery/tienda tokens",
             "Catches retail markets mislabeled as *_restaurant (Lama's Kitchen & "
             "Seafood Market, McManus Seafood, La Michoacana, etc.)."],
        ],
        widths=[1.4, 0.8, 1.5, 3.3],
    )
    doc.add_paragraph(
        "Net: 158 -> 104, and the obvious 'that's not a restaurant' rows a homebuyer "
        "would notice are largely gone."
    )

    # ---- 3. Trade-offs ----
    doc.add_heading("3. Known trade-offs of v3 (accepted for now)", level=1)
    for t in [
        "Min confidence 0.7 drops some real restaurants. Genuine spots tagged low in "
        "Overture (e.g. Festival Italiano 0.08, ZENSHI 0.27, Kame 0.30) will not appear. "
        "Intentional precision-over-recall choice; revisit after the freshness check (P1).",
        "Name-token market excludes are blunt. *market* would also drop a legitimate "
        "'Market Street Grill'-type name. None existed in the test set, but it is a "
        "nationwide risk - the durable fix is the is_retail_food_market flag (P2).",
        "Address dedup can hide a real co-located restaurant. ON is right here, but in "
        "food-hall settings it keeps only the highest-confidence tenant. Durable fix is "
        "dedup by (address + name), not address alone (P4).",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ---- 4. Not fixed ----
    doc.add_heading("4. Issues v3 does NOT fix (why the dev section exists)", level=1)
    doc.add_paragraph(
        "Even at 104 results, these slip through because the filter only reads Overture "
        "fields, and Overture is wrong or stale:"
    )
    grid_table(
        doc,
        ["Problem", "Example still in results", "Root cause"],
        [
            ["Permanently closed, tagged open",
             "Norris's Ribs (demolished 2017); conf 0.77 so it clears the gate",
             "operating_status is stale; no freshness check"],
            ["National fast-food chains",
             "Pizza Hut (brand_wikidata Q191615), Jack in the Box",
             "Tagged pizza_restaurant / burger_restaurant, not fast_food_restaurant"],
            ["Gas-station / counter brands",
             "Krispy Krunchy Chicken (no address)",
             "Tagged chicken_restaurant"],
            ["Same brand, multiple addresses",
             "Goodfella's Pizza x3, A Touch of Brooklyn x2, Sonny's BBQ x2, Dale's BBQ x2",
             "Address dedup cannot merge across distinct addresses / spellings"],
            ["Clubs / venues serving food",
             "American Legion of Ft. Pierce (alt bar)",
             "Tagged basic_category = restaurant"],
        ],
        widths=[1.7, 2.6, 2.4],
    )

    # ---- 5. Dev team ----
    doc.add_heading("5. Dev team: what to improve next", level=1)
    doc.add_paragraph(
        "The admin filter is good enough to ship. The remaining wins are data-layer, "
        "not filter-tuning - the same enrichment pattern used for the Hospital filter. "
        "In priority order:"
    )

    def dev_item(heading, problem, build, hook):
        doc.add_heading(heading, level=2)
        p = doc.add_paragraph()
        p.add_run("Problem: ").bold = True
        p.add_run(problem)
        p = doc.add_paragraph()
        p.add_run("Build: ").bold = True
        p.add_run(build)
        p = doc.add_paragraph()
        p.add_run("Payoff / filter hook: ").bold = True
        p.add_run(hook)

    dev_item(
        "P1 - Operating-status freshness (fixes 'closed but listed')",
        "Overture operating_status is stale; permanently-closed places (Norris's Ribs, "
        "closed 8 years) still read open. No boolean filter can catch this.",
        "Add status_verified_at (timestamp) and status_source (text) columns. A "
        "scheduled job re-confirms operating_status against a fresher source - Google "
        "Places Details (business_status) or OSM - for POIs shown to users. Drop/hide "
        "anything not confirmed open within N months (suggest 12).",
        "Add an 'Operating status verified within' control, or have the widget query "
        "exclude status_verified_at older than the interval.",
    )
    dev_item(
        "P2 - Retail-food-market flag (fixes markets mislabeled as restaurants)",
        "Retail seafood/meat/grocery markets are tagged *_restaurant (Pelican Seafood, "
        "McManus Seafood, La Michoacana). Today we catch them with blunt *market* name "
        "tokens, which risks dropping real 'Market Grill' names nationwide.",
        "Add is_retail_food_market (bool), populated by enrichment using category + name "
        "signals (market, grocery, butcher, fishmonger, tienda, meat market, seafood "
        "market) AND a secondary source (Google types contains grocery_or_supermarket / "
        "food without restaurant).",
        "Exclude is_retail_food_market = true, then remove the *market* name tokens from "
        "the admin filter. Also feeds the future Grocery filter.",
    )
    dev_item(
        "P3 - Chain / fast-food normalization (fixes Pizza Hut, Jack in the Box)",
        "National QSR chains are tagged as cuisine restaurants (pizza_restaurant, "
        "burger_restaurant), so category excludes miss them.",
        "Maintain a brand-based QSR exclude keyed on brand_wikidata / brand_name_primary "
        "(Pizza Hut Q191615, Jack in the Box, KFC, Domino's, etc.) - a small national "
        "reference list, not per-city. Add is_chain_fast_food (bool).",
        "Exclude is_chain_fast_food = true. Keep it a flag so it is easy to toggle per "
        "category/product.",
    )
    dev_item(
        "P4 - Smarter dedup (fixes same-brand duplicates & co-located places)",
        "Address-only dedup both misses same-brand duplicates at different addresses / "
        "spelling variants, and can hide a legitimately co-located second restaurant.",
        "Dedup on a composite key - normalized (name + address) plus a small geospatial "
        "tolerance (~25 m) - instead of address alone. Prefer the highest-confidence, "
        "most-complete record; keep distinct names at the same address.",
        "Replaces the current 'Deduplicate addresses' checkbox behavior.",
    )
    dev_item(
        "P5 - Confidence is not freshness (process note)",
        "Overture confidence measures tagging certainty, not whether a place is open or "
        "correctly categorized.",
        "Document this for the team. Once P1-P3 exist, lower min_confidence back toward "
        "~0.3 to recover the real low-confidence restaurants the 0.7 gate hides.",
        "Recovers recall lost to the current precision-first floor.",
    )

    doc.add_heading("Suggested column summary", level=2)
    grid_table(
        doc,
        ["Column", "Type", "Populated by", "Used to"],
        [
            ["status_verified_at", "timestamp", "P1 cron (Google/OSM)", "drop stale/closed"],
            ["status_source", "text", "P1 cron", "auditing"],
            ["is_retail_food_market", "bool", "P2 enrichment", "exclude markets / feed Grocery"],
            ["is_chain_fast_food", "bool", "P3 brand list", "exclude national QSR"],
        ],
        widths=[1.9, 1.0, 1.9, 1.9],
    )
    doc.add_paragraph(
        "Once P1-P3 ship, the admin filter can be simplified back to "
        "basic_category_include and category_primary_exclude and "
        "category_alternate_exclude plus the three new boolean excludes - and the "
        "name-token hacks and the high confidence floor can be retired."
    )

    # ---- Appendix ----
    doc.add_heading("Appendix A - config history", level=1)
    for t in [
        "v1: broad category includes (restaurant, casual_eatery, fine_dining in primary "
        "AND alternate). Leaked ~12 non-restaurants because many non-restaurants list "
        "*_restaurant as an alternate.",
        "v2 / v2.1: gated on basic_category = restaurant, cleared the primary/alternate "
        "includes, trimmed cuisine tokens out of the excludes. Dedup OFF, confidence 0.0. "
        "Result ~158 with retail markets and stale/closed places still present.",
        "v3 (current): confidence 0.7, dedup ON, added market/grocery/tienda name "
        "excludes. Result 104, shippable; residual issues are data-layer (see section 5).",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
