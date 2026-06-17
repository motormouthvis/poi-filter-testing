"""Build a formatted .docx for Google Docs (Upload -> Open with Google Docs)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_kv_table(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] = ("Field", "Value")):
    t = doc.add_table(rows=1 + len(rows), cols=2)
    t.style = "Table Grid"
    h = t.rows[0].cells
    h[0].text, h[1].text = header
    for i, (k, v) in enumerate(rows, start=1):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v


def add_one_col_table(doc: Document, title: str, lines: list[str]):
    doc.add_heading(title, level=3)
    t = doc.add_table(rows=len(lines), cols=1)
    t.style = "Table Grid"
    for i, line in enumerate(lines):
        t.rows[i].cells[0].text = line


def monospace_para(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    p.paragraph_format.keep_together = True


def main():
    out = Path(__file__).resolve().parent / "POI_Filter_Testing_Formatted.docx"

    NAME_EXCLUDES = [
        "veterinary",
        "veterinarian",
        "animal hospital",
        "pet hospital",
        "urgent care",
        "immediate care",
        "express care",
        "walk-in clinic",
        "walk in clinic",
        "minute clinic",
        "convenient care",
        "retail clinic",
        "telehealth",
        "dental",
        "dentist",
        "orthodont",
        "endodont",
        "oral surgery",
        "chiropract",
        "acupuncture",
        "med spa",
        "medspa",
        "plastic surgery",
        "cosmetic surgery",
        "bariatric center",
        "weight loss center",
        "fertility",
        "ivf",
        "dialysis",
        "sleep center",
        "sleep medicine",
        "insomnia",
        "pain management clinic",
        "imaging center",
        "radiology center",
        "mri",
        "ct scan",
        "laboratory",
        "blood draw",
        "plasma",
        "blood donation",
        "surgery center",
        "surgical center",
        "ambulatory surgery",
        "outpatient surgery",
        "endoscopy center",
        "chemotherapy",
        "radiation oncology",
        "cancer center",
        "oncology center",
        "cardiology office",
        "orthopedic clinic",
        "ent clinic",
        "eye clinic",
        "vision center",
        "hearing aid",
        "audiology",
        "physical therapy",
        "rehabilitation",
        "substance abuse",
        "detox",
        "psychiat",
        "mental health clinic",
        "counseling center",
        "hospice",
        "nursing home",
        "skilled nursing",
        "assisted living",
        "memory care",
        "long term care",
        "home health",
        "pharmacy",
        "drugstore",
        "surgical specialists",
        "suite ",  # trailing space intentional for some matchers
        "healthcare specialist",
        "trauma surgeons",
        "physicians group",
        "physician group",
    ]

    ALT_EXCLUDES = [
        "urgent_care",
        "walk_in",
        "walk-in",
        "clinic",
        "doctor",
        "health_clinic",
        "family_practice",
        "primary_care",
        "pediatric_clinic",
        "dentist",
        "dental",
        "orthodont",
        "chiropract",
        "acupuncture",
        "tattoo_and_piercing",
        "dog_park",
        "event_planning",
        "charity_organization",
        "community_services_non_profits",
        "sleep_specialist",
        "cancer_treatment_center",
        "dialysis_center",
        "imaging_center",
        "radiology_center",
        "laboratory",
        "laboratory_testing",
        "outpatient",
        "ambulatory",
        "hospice",
        "nursing",
        "assisted_living",
        "rehab",
        "rehabilitation",
        "mental_health",
        "counseling",
        "therapy",
        "surgical_center",
    ]

    doc = Document()
    t = doc.add_heading("POI Filter Testing and Results", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Use this file: upload to Google Drive, then Open with → Google Docs. "
        "Formatting (headings + tables) is preserved better than pasting plain text exports."
    )

    doc.add_heading("1. Hospital (USA) — Django admin preset", level=1)

    doc.add_heading("1.1 Saved filter / search fields (per query)", level=2)
    add_kv_table(
        doc,
        [
            ("Saved filter name", "Hospital (or your label)"),
            ("Latitude, Longitude", "User search pin — not part of static preset"),
            ("Distance (miles)", "Product default, e.g. 30"),
            ("Max results", "e.g. 250"),
            ("Min confidence", "Leave empty unless tuning (0–1)"),
            ("Address, City, State, Zipcode, Operating status", "(empty)"),
            ("Website", "Any"),
        ],
    )
    doc.add_paragraph()

    doc.add_heading("1.2 Business name — include (primary and common)", level=2)
    doc.add_paragraph("(empty in both boxes)")
    doc.add_paragraph()

    doc.add_heading("1.3 Business name — exclude (same list in primary and common)", level=2)
    add_one_col_table(doc, "Terms — one per line in both name exclude boxes", NAME_EXCLUDES)
    doc.add_paragraph()

    doc.add_heading("1.4 Basic category", level=2)
    add_kv_table(doc, [("Include", "hospital"), ("Exclude", "(empty)")])
    doc.add_paragraph()

    doc.add_heading("1.5 Category primary", level=2)
    doc.add_paragraph("Include (one per line):")
    add_one_col_table(doc, "Category primary — include (one per line)", ["hospital", "emergency_room", "medical_center"])
    doc.add_paragraph("Exclude: (empty)")
    doc.add_paragraph()

    doc.add_heading("1.6 Category alternate", level=2)
    doc.add_paragraph("Include (one per line):")
    add_one_col_table(doc, "Category alternate — include (one per line)", ["emergency_room", "medical_center"])
    doc.add_paragraph("Exclude (one per line):")
    add_one_col_table(doc, "Category alternate — exclude (one per line)", ALT_EXCLUDES)
    doc.add_paragraph()

    doc.add_heading("1.7 Taxonomy blocks", level=2)
    doc.add_paragraph("All four taxonomy boxes: (empty)")
    doc.add_paragraph()

    doc.add_heading("1.8 Query builder (final expression)", level=2)
    monospace_para(
        doc,
        "name_primary_exclude and name_common_exclude and basic_category_include and "
        "category_primary_include and category_alternate_include and category_alternate_exclude",
    )
    doc.add_paragraph(
        "If “Generate Query” renames tokens, keep your app’s official token names and preserve this boolean structure."
    )

    doc.add_heading("1.9 Known limitation", level=2)
    doc.add_paragraph(
        "Rows like Cleveland Clinic Martin South Hospital may have empty category_alternate and will not "
        "match category_alternate_include until you add something like acute_hospital_er_capable (cron / external verify)."
    )

    doc.add_heading("2. Dev brief — dedupe, under-tags, columns, evaluation", level=1)

    doc.add_heading("2.1 Goals", level=2)
    for line in [
        "Dedupe: one user-facing pin per same physical site (not “same health system”).",
        "Under-tags: include real acute hospitals when Overture has category_alternate empty.",
        "Nationwide: avoid city/name whitelists; prefer DB-owned fields + optional external verification.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2.2 Dedupe rules", level=2)
    for line in [
        "Merge only if normalized street address matches OR distance ≤ 50–100 m and same brand/operator if available.",
        "Do not merge different street addresses (main hospital vs freestanding ER elsewhere).",
        "Do merge same-campus duplicates (Hospital + Emergency Department at same address/coords).",
        "Output: canonical_poi_id or duplicate_of_id.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2.3 New DB columns (minimum)", level=2)
    t2 = doc.add_table(rows=5, cols=3)
    t2.style = "Table Grid"
    hdr = t2.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Column", "Type", "Purpose"
    data = [
        ("acute_hospital_er_capable", "boolean", "Filterable ER / acute hospital layer"),
        ("acute_hospital_source", "enum", "unknown | overture_tagged | inferred_internal | verified_external | manual_override"),
        ("acute_hospital_updated_at", "timestamp", "Audit"),
        ("acute_hospital_evidence (optional)", "json/text", "CMS id, OSM id, rule id"),
    ]
    for i, (a, b, c) in enumerate(data, start=1):
        t2.rows[i].cells[0].text = a
        t2.rows[i].cells[1].text = b
        t2.rows[i].cells[2].text = c
    doc.add_paragraph(
        "Filter change: keep current Overture logic OR acute_hospital_er_capable = true (shape depends on query builder)."
    )

    doc.add_heading("2.4 Population order", level=2)
    monospace_para(doc, "manual_override > verified_external > overture_tagged > inferred_internal > else false/unknown")

    doc.add_heading("2.5 Benchmark (Treasure Coast test pin)", level=2)
    tb = doc.add_table(rows=1 + 9, cols=4)
    tb.style = "Table Grid"
    bh = tb.rows[0].cells
    bh[0].text, bh[1].text, bh[2].text, bh[3].text = "#", "Reference", "Overture / filter", "Action"
    bench = [
        ("1", "HCA Lawnwood", "Passes strict tags", "overture_tagged"),
        ("2", "Indian River", "Passes strict tags", "overture_tagged"),
        ("3", "Florida Coast (new)", "Often missing in snapshot", "refresh / external"),
        ("4", "St Lucie (Tiffany)", "Passes (e.g. St Lucie Medical Center)", "overture_tagged"),
        ("4b", "Darwin Square ER", "Separate POI; different address", "product / keep"),
        ("5", "Tradition + ED", "Same site → two rows", "dedupe"),
        ("6", "Martin North", "Passes", "overture_tagged"),
        ("7", "Sebastian River (13695 US-1)", "Passes", "overture_tagged"),
        ("8", "Martin South (Salerno)", "Empty alternate → strict filter drops", "verified_external or inferred_internal"),
    ]
    for i, row in enumerate(bench, start=1):
        for j, val in enumerate(row):
            tb.rows[i].cells[j].text = val

    doc.add_heading("3. How to get this into Google Docs", level=1)
    for line in [
        "Google Drive → New → File upload → select POI_Filter_Testing_Formatted.docx",
        "Right-click the uploaded file → Open with → Google Docs (Google converts it; you get an editable Doc).",
        "Or: In Google Docs → File → Open → Upload the .docx.",
    ]:
        doc.add_paragraph(line, style="List Number")

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
